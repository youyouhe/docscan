#!/usr/bin/env python3
"""
DocScan style-transfer engine — copy paragraph-style definitions (headings +
body) from a sample .docx onto a target .docx so the target adopts the sample's
visual look, while leaving the target's content untouched.

Pairing is role-based, NOT by styleId/name equality:
  * Headings pair by outlineLvl (0..N) — the only stable cross-document anchor.
    Sample "heading 1" (styleId "2", outlineLvl 0) pairs with target "Heading1"
    (outlineLvl 0) regardless of id/naming drift.
  * Body pairs by semantic name (Normal / Body Text / 段落正文 ...).

Style definitions are carried by deep-copying the sample's <w:pPr>/<w:rPr> into
the target style element (replacing any existing block, same position). Theme
fonts (majorFont/minorFont in theme1.xml) are synced so themeFont references
(majorHAnsi / minorEastAsia) resolve identically.

Caller owns the on-disk files: apply_sample_styles(target_path, sample_path)
loads both, mutates target in place, saves it, and returns an applied report.
"""

import copy
import os
import re
import zipfile

from docx import Document
from docx.oxml.ns import qn

# Body-text style names (canonical + common Chinese aliases). OutlineLvl-less
# paragraph styles whose name matches one of these are treated as the "body"
# role. Matched case-insensitively after stripping whitespace.
_BODY_NAMES = {
    'normal', 'body text', '段落正文', '正文', '论文正文', '正文文本',
}

# Matches "heading 3" / "标题3" / "Heading 12" → captures the trailing number.
_HEADING_NUM_RE = re.compile(r'(?:heading|标题)\s*([1-9])', re.IGNORECASE)

# DrawingML main namespace (a:) — used for theme1.xml fontScheme surgery.
_A = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
_THEME_ENTRY = 'word/theme/theme1.xml'


# ════════════════════════════════════════════════════════════════════
#  Style introspection
# ════════════════════════════════════════════════════════════════════
def _styles_root(doc):
    """The <w:styles> root element of a python-docx Document."""
    return doc.styles.element


def _style_info(st_elem):
    """Parse one <w:style> element into a plain dict:
        {id, name, type, outlineLvl (int or None), elem}
    `elem` is kept so callers can reach back to the live OOXML node.
    """
    name_el = st_elem.find(qn('w:name'))
    pPr = st_elem.find(qn('w:pPr'))
    outline = None
    if pPr is not None:
        ol = pPr.find(qn('w:outlineLvl'))
        if ol is not None:
            try:
                outline = int(ol.get(qn('w:val')))
            except (TypeError, ValueError):
                outline = None
    return dict(
        id=st_elem.get(qn('w:styleId')) or '',
        name=(name_el.get(qn('w:val')) if name_el is not None else ''),
        type=(st_elem.get(qn('w:type')) or ''),
        outlineLvl=outline,
        elem=st_elem,
    )


def _paragraph_styles(doc):
    """List of _style_info for every paragraph-type style in doc."""
    return [info for info in (_style_info(st) for st in _styles_root(doc).findall(qn('w:style')))
            if info['type'] == 'paragraph']


def _heading_num(name):
    m = _HEADING_NUM_RE.search(name or '')
    return int(m.group(1)) if m else None


# ════════════════════════════════════════════════════════════════════
#  Role pairing
# ════════════════════════════════════════════════════════════════════
def map_roles(sample_styles, target_styles):
    """Pair target paragraph styles to sample paragraph styles by role.

    Returns (heading_pairs, body_pairs) where each pair is
    (target_styleId, sample_styleId).

    Headings: grouped by outlineLvl. Within a level, prefer the sample style
    whose name encodes the same heading number as the target; failing that the
    one whose name encodes (outlineLvl+1); else the first at that level.

    Body: target styles whose name is a body alias AND have no outlineLvl,
    paired with the sample body of the same name, else the sample "Normal".
    """
    def by_outline(styles):
        d = {}
        for s in styles:
            ol = s['outlineLvl']
            if ol is not None:
                d.setdefault(ol, []).append(s)
        return d

    s_ol = by_outline(sample_styles)
    t_ol = by_outline(target_styles)

    heading_pairs = []
    for ol, tlist in sorted(t_ol.items()):
        slist = s_ol.get(ol)
        if not slist:
            continue
        # target's own heading number (e.g. "Heading 2" → 2); else level+1.
        t_num = _heading_num(tlist[0]['name']) if tlist else None
        chosen = None
        want = t_num if t_num is not None else ol + 1
        chosen = next((s for s in slist if _heading_num(s['name']) == want), None)
        if chosen is None:
            chosen = slist[0]
        heading_pairs.append((tlist[0]['id'], chosen['id']))

    s_body = [s for s in sample_styles
              if s['outlineLvl'] is None and (s['name'] or '').strip().lower() in _BODY_NAMES]
    body_pairs = []
    for t in target_styles:
        if t['outlineLvl'] is not None:
            continue
        if (t['name'] or '').strip().lower() not in _BODY_NAMES:
            continue
        same = next((s for s in s_body
                     if s['name'].strip().lower() == t['name'].strip().lower()), None)
        chosen = same or next((s for s in s_body
                               if (s['name'] or '').strip().lower() == 'normal'), None)
        if chosen:
            body_pairs.append((t['id'], chosen['id']))

    return heading_pairs, body_pairs


# ════════════════════════════════════════════════════════════════════
#  Definition transfer
# ════════════════════════════════════════════════════════════════════
def _replace_block(parent, tag_qn, sample_block):
    """Replace parent's first child named tag_qn with a deep copy of
    sample_block; if absent, append it. Position is preserved on replace."""
    old = parent.find(tag_qn)
    cloned = copy.deepcopy(sample_block)
    if old is not None:
        parent.replace(old, cloned)
    else:
        parent.append(cloned)


def _find_representative_paragraph(doc, style_id):
    """First paragraph in the sample using pStyle=style_id (body + tables).
    Returns the <w:p> element, or None if the style is defined but unused."""
    for p in doc.element.body.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue
        ps = pPr.find(qn('w:pStyle'))
        if ps is not None and ps.get(qn('w:val')) == style_id:
            return p
    return None


def _first_run_rpr(para):
    """The <w:rPr> of the first run in `para` that carries one, else None."""
    for r in para.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            return rPr
    return None


def _merge_props(base_block, overlay_block, exclude=()):
    """Deep-merge overlay's children into a copy of base: same-tag child
    replaces, new tags append. OOXML pPr/rPr children are unique by tag, so
    dict-by-tag semantics hold. `exclude` drops named tags from the overlay
    (used to keep pStyle out of a style definition)."""
    if base_block is None and overlay_block is None:
        return None
    if base_block is None:
        return copy.deepcopy(overlay_block)
    result = copy.deepcopy(base_block)
    if overlay_block is None:
        return result
    for child in overlay_block:
        if child.tag in exclude:
            continue
        old = result.find(child.tag)
        if old is not None:
            result.replace(old, copy.deepcopy(child))
        else:
            result.append(copy.deepcopy(child))
    return result


def _effective_blocks(sample_doc, sample_elem):
    """The pPr/rPr a sample style *renders with* = the style's own definition
    merged with the direct formatting on a representative paragraph.

    Samples often set the real look via direct run formatting (selecting a
    heading and changing its size in the toolbar) rather than by editing the
    style definition — so the definition alone (sz=44) doesn't match what the
    user sees (sz=32). Folding in the representative paragraph's direct
    formatting captures the effective look. pStyle itself is excluded so we
    don't write a self-reference into the target style definition."""
    s_pPr = sample_elem.find(qn('w:pPr'))
    s_rPr = sample_elem.find(qn('w:rPr'))
    sid = sample_elem.get(qn('w:styleId'))
    para = _find_representative_paragraph(sample_doc, sid)
    if para is None:
        return s_pPr, s_rPr
    d_pPr = para.find(qn('w:pPr'))
    d_rPr = _first_run_rpr(para)
    # Exclude pStyle (self-reference), rPr (paragraph-mark character format),
    # and numPr (list numbering) from the paragraph-direct overlay. Numbering
    # is handled by a dedicated sync step (sync_numbering) that also carries
    # the numbering.xml definitions; merging numPr here would write a sample
    # numId that has no definition in the target.
    eff_pPr = (_merge_props(s_pPr, d_pPr,
                            exclude=(qn('w:pStyle'), qn('w:rPr'), qn('w:numPr')))
               if d_pPr is not None else s_pPr)
    eff_rPr = (_merge_props(s_rPr, d_rPr) if d_rPr is not None else s_rPr)
    return eff_pPr, eff_rPr

def _apply_pair(target_elem, sample_doc, sample_elem):
    """Overwrite target style's formatting with the sample style's EFFECTIVE
    formatting (definition + direct formatting).

    pPr is replaced wholesale (spacing/alignment/indent must come entirely
    from the sample). rPr is MERGED onto the target's existing rPr so that
    attributes the sample inherits via basedOn but doesn't define inline
    (notably <w:rFonts>) are preserved — replacing rPr wholesale would drop
    the target's theme-font binding, fall back to docDefaults, and lose bold
    rendering for CJK fonts that lack a bold variant."""
    eff_pPr, eff_rPr = _effective_blocks(sample_doc, sample_elem)
    if eff_pPr is not None:
        _replace_block(target_elem, qn('w:pPr'), eff_pPr)
    if eff_rPr is not None:
        t_rPr = target_elem.find(qn('w:rPr'))
        _replace_block(target_elem, qn('w:rPr'), _merge_props(t_rPr, eff_rPr))


def apply_style_definitions(target_doc, sample_doc, heading_pairs, body_pairs):
    """Mutate target_doc's styles in place. Returns a list of applied records
    [{role, targetId, sampleId, sampleName}]."""
    s_by_id = {st.get(qn('w:styleId')): st
               for st in _styles_root(sample_doc).findall(qn('w:style'))}
    t_by_id = {st.get(qn('w:styleId')): st
               for st in _styles_root(target_doc).findall(qn('w:style'))}
    applied = []
    for pairs, role in ((heading_pairs, 'heading'), (body_pairs, 'body')):
        for t_id, s_id in pairs:
            t_el = t_by_id.get(t_id)
            s_el = s_by_id.get(s_id)
            if t_el is None or s_el is None:
                continue
            _apply_pair(t_el, sample_doc, s_el)
            s_name_el = s_el.find(qn('w:name'))
            s_name = s_name_el.get(qn('w:val')) if s_name_el is not None else ''
            applied.append(dict(role=role, targetId=t_id,
                                sampleId=s_id, sampleName=s_name))
    return applied


# ════════════════════════════════════════════════════════════════════
#  Document-defaults font baseline sync (docDefaults/rPrDefault/rPr/rFonts)
# ════════════════════════════════════════════════════════════════════
def sync_doc_defaults(target_doc, sample_doc):
    """Carry the sample's document-level font baseline — the <w:rFonts> in
    docDefaults/rPrDefault/rPr — onto the target, so every text that inherits
    the default resolves to the same ascii/eastAsia/hAnsi/cs font (or theme
    slot) the sample uses. Only rFonts is touched; baseline size/lang stay.

    Without this, themeFont slots are synced (sync_theme_fonts) but the
    *mapping* of a paragraph to a slot (asciiTheme=minorHAnsi etc.) can still
    differ, so a sample that pins eastAsia to a named font (e.g. 黑体) would
    lose it on the target. Returns True if synced."""
    s_dd = _styles_root(sample_doc).find(qn('w:docDefaults'))
    t_dd = _styles_root(target_doc).find(qn('w:docDefaults'))
    if s_dd is None or t_dd is None:
        return False
    s_rpr = s_dd.find(f'{qn("w:rPrDefault")}/{qn("w:rPr")}')
    if s_rpr is None:
        return False
    s_rfonts = s_rpr.find(qn('w:rFonts'))
    if s_rfonts is None:
        return False
    # Ensure target has rPrDefault/rPr (python-docx always does, but be safe).
    from docx.oxml import OxmlElement
    t_rprd = t_dd.find(qn('w:rPrDefault'))
    if t_rprd is None:
        t_rprd = OxmlElement('w:rPrDefault')
        t_dd.insert(0, t_rprd)
    t_rpr = t_rprd.find(qn('w:rPr'))
    if t_rpr is None:
        t_rpr = OxmlElement('w:rPr')
        t_rprd.insert(0, t_rpr)
    cloned = copy.deepcopy(s_rfonts)
    old = t_rpr.find(qn('w:rFonts'))
    if old is not None:
        t_rpr.replace(old, cloned)     # keeps position (rFonts must stay first)
    else:
        t_rpr.insert(0, cloned)        # rFonts is schema-required first child
    return True


# ════════════════════════════════════════════════════════════════════
#  Theme font sync (majorFont / minorFont)
# ════════════════════════════════════════════════════════════════════
def _rewrite_zip_entry(zip_path, entry_name, new_bytes):
    """Replace a single entry inside a .docx zip in place, preserving all
    other entries and their compression."""
    from lxml import etree  # local import; only needed on this path
    tmp = zip_path.with_suffix(zip_path.suffix + '.tmp')
    with zipfile.ZipFile(zip_path, 'r') as zin, \
         zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == entry_name:
                zout.writestr(item, new_bytes)
            else:
                zout.writestr(item, zin.read(item.filename))
    os.replace(tmp, zip_path)


def sync_theme_fonts(target_path, sample_path):
    """Copy majorFont/minorFont from sample's theme1.xml into target's, so
    themeFont references (majorHAnsi / minorEastAsia) resolve identically.

    Only the two font child elements are swapped — theme colors and other
    theme parts are left untouched. Returns True if synced, False if either
    side had no usable fontScheme (then nothing is changed)."""
    from lxml import etree

    try:
        s_xml = zipfile.ZipFile(str(sample_path)).read(_THEME_ENTRY)
    except (KeyError, FileNotFoundError):
        return False
    s_root = etree.fromstring(s_xml)
    s_fs = s_root.find(f'.//{_A}fontScheme')
    if s_fs is None:
        return False
    s_major = s_fs.find(f'{_A}majorFont')
    s_minor = s_fs.find(f'{_A}minorFont')
    if s_major is None and s_minor is None:
        return False

    with zipfile.ZipFile(str(target_path)) as z:
        names = z.namelist()
        if _THEME_ENTRY not in names:
            return False
        t_xml = z.read(_THEME_ENTRY)

    t_root = etree.fromstring(t_xml)
    t_fs = t_root.find(f'.//{_A}fontScheme')
    if t_fs is None:
        return False
    if s_major is not None:
        t_old = t_fs.find(f'{_A}majorFont')
        if t_old is not None:
            t_fs.replace(t_old, copy.deepcopy(s_major))
    if s_minor is not None:
        t_old = t_fs.find(f'{_A}minorFont')
        if t_old is not None:
            t_fs.replace(t_old, copy.deepcopy(s_minor))
    new_xml = etree.tostring(t_root, xml_declaration=True,
                             encoding='UTF-8', standalone=True)
    _rewrite_zip_entry(target_path, _THEME_ENTRY, new_xml)
    return True


# ════════════════════════════════════════════════════════════════════
#  Heading-numbering sync (numbering.xml definitions + style-level numPr)
# ════════════════════════════════════════════════════════════════════
_NUMBERING_ENTRY = 'word/numbering.xml'
_MAX_HEADING_LVL = 9


def _heading_numpr_map(sample_doc):
    """Per heading outlineLvl, return the (numId, ilvl) to use for numbering,
    or omit the level if it shouldn't be auto-numbered.

    Decision rule: the FIRST paragraph of each outline level decides. The
    first occurrence of a heading level sets that level's standard look — if
    it carries <w:numPr> the level is numbering-driven; if not (e.g. the text
    itself writes "第一章" / "一、") the level stays unnumbered. This avoids
    mis-promoting the mixed numbering found in real templates (where a few
    sub-clauses use numbering fields but the section heads use text numbers).

    Returns {outlineLvl: (numId_str, ilvl_str)}; empty when unnumbered."""
    numbered = {}      # ol -> (numId, ilvl), only for levels whose first para is numbered
    decided = set()    # outlineLvls whose first paragraph has been seen
    styles_by_id = {st.get(qn('w:styleId')): _style_info(st)
                    for st in _styles_root(sample_doc).findall(qn('w:style'))}
    for p in sample_doc.element.body.iter(qn('w:p')):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue
        ps = pPr.find(qn('w:pStyle'))
        if ps is None:
            continue
        info = styles_by_id.get(ps.get(qn('w:val')))
        if info is None or info['outlineLvl'] is None:
            continue
        ol = info['outlineLvl']
        if ol >= _MAX_HEADING_LVL or ol in decided:
            continue          # first occurrence already decided this level
        decided.add(ol)
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            continue          # first para unnumbered → level stays unnumbered
        numId = numPr.find(qn('w:numId'))
        ilvl = numPr.find(qn('w:ilvl'))
        if numId is not None and ilvl is not None:
            numbered[ol] = (numId.get(qn('w:val')), ilvl.get(qn('w:val')))
    # Multi-level numbering is only coherent when it starts at the top level
    # and is contiguous: if ol=0 isn't numbered, a numbered ol=1 would dangle
    # without a parent counter. Keep only the contiguous run from ol=0 upward.
    contiguous = {}
    for ol in range(_MAX_HEADING_LVL):
        if ol in numbered:
            contiguous[ol] = numbered[ol]
        else:
            break
    return contiguous


def sync_numbering(target_path, sample_path, heading_pairs):
    """If the sample's headings are numbered, carry the numbering definitions
    into the target's numbering.xml and bind each target heading style to the
    right (numId, ilvl) so all headings auto-number with the sample's format.

    The sample's paragraph-level numPr is promoted to a STYLE-level numPr on
    the target (cleaner than per-paragraph, and pandoc-generated targets have
    no paragraph numPr to begin with). abstractNum/num ids are remapped to
    avoid colliding with the target's existing numbering. Returns True if the
    target was modified (sample had heading numbering), False otherwise."""
    from lxml import etree
    sample_doc = Document(str(sample_path))
    lvl_map = _heading_numpr_map(sample_doc)
    if not lvl_map:                        # sample headings are unnumbered
        return False

    # Distinct sample numIds used by the headings, in first-seen order.
    sample_num_ids = []
    for ol in sorted(lvl_map):
        nid = lvl_map[ol][0]
        if nid not in sample_num_ids:
            sample_num_ids.append(nid)

    # ---- read both numbering.xml trees ----
    try:
        s_xml = zipfile.ZipFile(str(sample_path)).read(_NUMBERING_ENTRY)
    except (KeyError, FileNotFoundError):
        return False
    s_root = etree.fromstring(s_xml)
    s_abstract = {a.get(qn('w:abstractNumId')): a
                  for a in s_root.findall(qn('w:abstractNum'))}
    s_num_to_abstract = {}
    for n in s_root.findall(qn('w:num')):
        nid = n.get(qn('w:numId'))
        an = n.find(qn('w:abstractNumId'))
        if an is not None:
            s_num_to_abstract[nid] = an.get(qn('w:val'))

    with zipfile.ZipFile(str(target_path)) as z:
        has_numbering = _NUMBERING_ENTRY in z.namelist()
        t_xml = z.read(_NUMBERING_ENTRY) if has_numbering else None

    if t_xml is not None:
        t_root = etree.fromstring(t_xml)
    else:
        t_root = etree.fromstring(
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

    # ---- remap ids past the target's current maxima ----
    def _max_id(root, tag, attr):
        vals = [int(e.get(qn(attr)) or 0) for e in root.findall(qn(tag))]
        return max(vals) if vals else -1
    base_abstract = _max_id(t_root, 'w:abstractNum', 'w:abstractNumId') + 1
    base_num = _max_id(t_root, 'w:num', 'w:numId') + 1

    abs_remap, num_remap = {}, {}
    for i, snid in enumerate(sample_num_ids):
        said = s_num_to_abstract.get(snid)
        if said is None or said not in s_abstract:
            continue
        new_aid = str(base_abstract + i)
        new_nid = str(base_num + i)
        abs_remap[said] = new_aid
        num_remap[snid] = new_nid
        # clone the abstractNum with remapped id
        cloned_abs = copy.deepcopy(s_abstract[said])
        cloned_abs.set(qn('w:abstractNumId'), new_aid)
        t_root.append(cloned_abs)
        # clone the num pointing at it
        cloned_num = etree.SubElement(t_root, qn('w:num'))
        cloned_num.set(qn('w:numId'), new_nid)
        an_el = etree.SubElement(cloned_num, qn('w:abstractNumId'))
        an_el.set(qn('w:val'), new_aid)

    if not num_remap:
        return False

    new_xml = etree.tostring(t_root, xml_declaration=True,
                             encoding='UTF-8', standalone=True)
    _rewrite_zip_entry(target_path, _NUMBERING_ENTRY, new_xml)

    # ---- bind target heading styles (in-memory, re-open to edit styles) ----
    target_doc = Document(str(target_path))
    t_by_id = {st.get(qn('w:styleId')): st
               for st in _styles_root(target_doc).findall(qn('w:style'))}
    for t_id, _s_id in heading_pairs:
        t_el = t_by_id.get(t_id)
        if t_el is None:
            continue
        t_info = _style_info(t_el)
        if t_info['outlineLvl'] is None:
            continue
        ol = t_info['outlineLvl']
        if ol not in lvl_map:
            continue
        snid, ilvl = lvl_map[ol]
        new_nid = num_remap.get(snid)
        if new_nid is None:
            continue
        pPr = t_el.find(qn('w:pPr'))
        if pPr is None:
            from docx.oxml import OxmlElement
            pPr = OxmlElement('w:pPr')
            t_el.insert(0, pPr)
        old = pPr.find(qn('w:numPr'))
        if old is not None:
            pPr.remove(old)
        from docx.oxml import OxmlElement
        numPr = OxmlElement('w:numPr')
        ilvl_el = OxmlElement('w:ilvl'); ilvl_el.set(qn('w:val'), ilvl)
        numId_el = OxmlElement('w:numId'); numId_el.set(qn('w:val'), new_nid)
        numPr.append(ilvl_el); numPr.append(numId_el)
        pPr.append(numPr)
    target_doc.save(str(target_path))
    return True


# ════════════════════════════════════════════════════════════════════
#  Entry point
# ════════════════════════════════════════════════════════════════════
def apply_sample_styles(target_path, sample_path):
    """Transfer heading + body paragraph styles from sample onto target,
    rewriting target_path in place. Returns an applied-report dict:

        {applied: [{role, targetId, sampleId, sampleName}, ...],
         themeFontsSynced: bool,
         docDefaultsSynced: bool,
         numberingSynced: bool}
    """
    sample_doc = Document(str(sample_path))
    target_doc = Document(str(target_path))
    s_styles = _paragraph_styles(sample_doc)
    t_styles = _paragraph_styles(target_doc)
    heading_pairs, body_pairs = map_roles(s_styles, t_styles)
    applied = apply_style_definitions(target_doc, sample_doc,
                                      heading_pairs, body_pairs)
    defaults_ok = sync_doc_defaults(target_doc, sample_doc)   # font baseline (pre-save)
    target_doc.save(str(target_path))           # persists style + docDefaults edits
    theme_ok = sync_theme_fonts(target_path, sample_path)  # then patch theme1.xml
    numbering_ok = sync_numbering(target_path, sample_path, heading_pairs)  # heading numbers
    return dict(applied=applied, themeFontsSynced=theme_ok,
                docDefaultsSynced=defaults_ok, numberingSynced=numbering_ok)
