#!/usr/bin/env python3
"""
DocScan API  —  FastAPI service for docx → PDF + Markdown conversion.

Quick start:
    python3 api.py
    → API at http://localhost:8800/api/
    → Swagger at http://localhost:8800/api/docs
    → Frontend demo at http://localhost:8800/

Endpoints:
    POST /api/convert              upload .docx → returns {id, totalPages, pdfUrl, mdUrl}
    GET  /api/pdf/{id}             download the generated PDF
    GET  /api/md/{id}              all-page Markdown array
    GET  /api/md/{id}/{page}       single-page Markdown
    GET  /api/conversions          list recent conversions
    GET  /api/health               health check

    POST /api/md2docx                    upload .md → returns {id, fileName, docxUrl}
    GET  /api/docx/{id}                  download the current docx
    GET  /api/docx/{id}/placeholders      list 【...】 placeholders with stable ids
    POST /api/docx/{id}/replace           replace placeholders by id
    GET  /api/docx/{id}/tables            list table structure (for crossref target picking)
    GET  /api/docx/{id}/preview           full-text preview (body paragraphs + tables)
    POST /api/docx/{id}/crossref          bookmark a body keyword + insert a page-number
                                           cross-reference field into a target table cell

The frontend ONLYOFFICE viewer is proxied through this server (same-origin
at /oo/…) so the demo at / also works.
"""

import asyncio, json, logging, os, re, secrets, shutil, subprocess, time, uuid, sys
from contextlib import asynccontextmanager
from pathlib import Path
from urllib.parse import urlparse, urlunparse

import httpx
import uvicorn
from docx import Document
from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# ——— reuse the conversion engine from server.py ———
ROOT = Path(__file__).parent.resolve()
sys.path.insert(0, str(ROOT))
from server import (
    _convert_docx_to_pdf,
    _extract_pdf_pages,
    _recalculate_fields_docx,
)
import docx_ops
import style_ops
OO_BACKEND = 'http://localhost:8079'   # ONLYOFFICE Docker container

log = logging.getLogger('docscan')
_FID_RE = re.compile(r'^[0-9a-f]{10,12}$')   # fid = uuid4().hex[:10] 或 [:12]

# ——— per-IP rate limits (sliding 60s window) & upload content guards ———
RATE_WINDOW = 60
RATE_READ = int(os.environ.get('DOCSCAN_RATE_READ', '60'))    # GET / IP / 分钟
RATE_WRITE = int(os.environ.get('DOCSCAN_RATE_WRITE', '10'))  # POST 等 / IP / 分钟
DOCX_MAX_UNCOMPRESSED = int(os.environ.get('DOCSCAN_DOCX_MAX_UNCOMPRESSED_MB', '200')) * 1024 * 1024

def _safe_path(base, fid, ext):
    """base/{fid}.{ext}，校验 fid 为十六进制以防路径遍历（公网加固）。"""
    if not _FID_RE.match(fid):
        raise HTTPException(404, 'not found')
    return base / f'{fid}.{ext}'

def _validate_docx(data: bytes) -> bytes:
    """Reject non-ZIP or zip-bomb .docx uploads."""
    if not data.startswith(b'PK\x03\x04'):
        raise HTTPException(400, 'not a valid .docx (ZIP) file')
    import io, zipfile
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
        total = sum(zi.file_size for zi in zf.infolist())
        n = len(zf.namelist())
    except Exception:
        raise HTTPException(400, 'corrupt .docx archive')
    if total > DOCX_MAX_UNCOMPRESSED:
        raise HTTPException(400, 'docx uncompressed size too large')
    if n > 10000:
        raise HTTPException(400, 'docx has too many entries')
    return data

def _validate_md(data: bytes) -> bytes:
    try:
        data.decode('utf-8')
    except UnicodeDecodeError:
        raise HTTPException(400, '.md file is not valid UTF-8')
    return data

_valid_keys_cache = set()
_valid_keys_mtime = -1.0
def _valid_keys():
    """有效 key 集合：DOCSCAN_API_KEY + .docscan-api-keys 每行（# 注释/空行跳过）。
    按 mtime 缓存，文件改动后自动刷新——增删 key 无需重启。"""
    global _valid_keys_cache, _valid_keys_mtime
    p = ROOT / '.docscan-api-keys'
    try:
        m = p.stat().st_mtime
    except OSError:
        return {API_KEY}
    if m != _valid_keys_mtime:
        keys = {API_KEY}
        try:
            for line in p.read_text('utf-8').splitlines():
                line = line.strip()
                if line and not line.startswith('#'):
                    keys.add(line)
        except OSError:
            pass
        _valid_keys_cache = keys
        _valid_keys_mtime = m
    return _valid_keys_cache

def _record_owner(fid, key):
    """上传时记录文档归属的 key（持久化到 owners/，重启不丢）。"""
    if key:
        (OWNERS_DIR / f'{fid}.owner').write_text(key, 'utf-8')

def _owner_of(fid):
    """fid 的归属 key；无记录返回 None。调用方负责先校验 fid 格式
    （见 _assert_owner / list_conv），本函数只负责读。"""
    p = OWNERS_DIR / f'{fid}.owner'
    try:
        return p.read_text('utf-8').strip() or None
    except OSError:
        return None

def _assert_owner(fid, request):
    """访问校验：仅 owner key 可操作该 fid；无归属记录(旧文件)→ 404。
    先把 fid 校验为十六进制，防止被用来路径穿越到 OWNERS_DIR 之外——
    _safe_path 也校验，但它在调用处排在 _assert_owner 之后，得各报各的。"""
    if not _FID_RE.match(fid):
        raise HTTPException(404, 'not found')
    key = getattr(request.state, 'api_key', None)
    if not key:
        return
    owner = _owner_of(fid)
    if owner is None:
        raise HTTPException(404, 'not found')
    if owner != key:
        raise HTTPException(403, 'forbidden')

# ——— data dirs ———
DOCS_DIR = ROOT / 'docs'
PDFS_DIR = ROOT / 'pdfs'
MDS_DIR  = ROOT / 'mds'
DOCX_DIR = ROOT / 'docx_store'   # persistent editable docx (md2docx output, placeholder/crossref edits)
OWNERS_DIR = ROOT / 'owners'     # {fid}.owner → key，文档归属（访问控制 S6）
for d in (DOCS_DIR, PDFS_DIR, MDS_DIR, DOCX_DIR, OWNERS_DIR):
    d.mkdir(parents=True, exist_ok=True)

conversions = {}   # {id: metadata}   in-memory
docx_docs = {}      # {id: metadata}   in-memory, tracks editable docx (see DOCX_DIR)

# 转换是同步阻塞调用（subprocess + httpx.Client），丢进线程池跑，
# 避免卡住 uvicorn 的单个事件循环；并发数上限防止把 ONLYOFFICE 转换 worker 打爆。
CONVERT_CONCURRENCY = int(os.environ.get('DOCSCAN_CONVERT_CONCURRENCY', '8'))
_convert_semaphore = asyncio.Semaphore(CONVERT_CONCURRENCY)
# 有界排队（背压）：执行中(CONVERT_CONCURRENCY) + 排队等待(MAX_QUEUED) 之外
# 的请求立即 503，避免突发流量下无限排队、把内存/连接撑爆。信号量只卡"同时
# 执行"，挡不住"无限排队"——这里给排队也设上限。
MAX_QUEUED = int(os.environ.get('DOCSCAN_MAX_QUEUED', '10'))
_concurrent_or_queued = 0   # 执行中 + 排队中的总数

@asynccontextmanager
async def _convert_slot():
    """占一个转换名额（正在执行或在信号量前排队），离开时归还。名额满
    （超过 并发 + 排队上限）直接 503，把"无限排队"变成"有界排队"。

    _concurrent_or_queued 的检查与自增之间无 await，在 asyncio 单线程里对
    协程天然原子，无需加锁；被 503 的请求在自增前就 raise，不占名额。"""
    global _concurrent_or_queued
    if _concurrent_or_queued >= CONVERT_CONCURRENCY + MAX_QUEUED:
        raise HTTPException(503, 'server busy: conversion queue full, retry later')
    _concurrent_or_queued += 1
    try:
        async with _convert_semaphore:   # 在此排队等执行名额
            yield
    finally:
        _concurrent_or_queued -= 1

# ——— upload size cap (stream-read so a huge body can't OOM us) ———
MAX_UPLOAD_BYTES = int(os.environ.get('DOCSCAN_MAX_UPLOAD_MB', '100')) * 1024 * 1024

# ——— retention: drop generated files older than N hours on startup (0 = off) ———
RETENTION_HOURS = int(os.environ.get('DOCSCAN_RETENTION_HOURS', '0'))

async def _read_capped(file):
    """Read an UploadFile in 1MB chunks, rejecting bodies over MAX_UPLOAD_BYTES
    (413) instead of slurping the whole thing into memory unbounded."""
    total, chunks = 0, []
    while True:
        chunk = await file.read(1024 * 1024)
        if not chunk:
            break
        total += len(chunk)
        if total > MAX_UPLOAD_BYTES:
            raise HTTPException(413, f'file too large (max {MAX_UPLOAD_BYTES // (1024 * 1024)}MB)')
        chunks.append(chunk)
    return b''.join(chunks)

def _cleanup_old_files():
    """Delete generated files older than RETENTION_HOURS. Off by default — only
    runs when an operator opts in via DOCSCAN_RETENTION_HOURS, so we never
    silently delete a user's data."""
    if RETENTION_HOURS <= 0:
        return
    cutoff = time.time() - RETENTION_HOURS * 3600
    for d in (PDFS_DIR, MDS_DIR, DOCX_DIR, DOCS_DIR):
        for p in d.glob('*'):
            try:
                if p.is_file() and p.stat().st_mtime < cutoff:
                    p.unlink()
            except OSError:
                pass

_cleanup_old_files()

# ——— API key: required for all /api/* except health/docs. If none is ——————
# ——— configured we mint an ephemeral one and print it, so the server never —
# ——— boots unauthenticated by accident. ————————————————————————————————
API_KEY = os.environ.get('DOCSCAN_API_KEY') or secrets.token_urlsafe(24)
# Comma-separated allow-list of browser origins allowed to call /api cross-site.
# Default empty = same-origin only (safest for public deploy). Set e.g.
# DOCSCAN_CORS_ORIGINS=https://app.example.com to open specific origins.
CORS_ORIGINS = [o.strip() for o in os.environ.get('DOCSCAN_CORS_ORIGINS', '').split(',') if o.strip()]
if not os.environ.get('DOCSCAN_API_KEY'):
    sys.stderr.write(
        '\n==============================================================\n'
        '  No DOCSCAN_API_KEY set — generated ephemeral key:\n'
        f'      {API_KEY}\n'
        '  Set DOCSCAN_API_KEY to keep it stable across restarts.\n'
        '==============================================================\n\n'
    )

# ═══════════════════════════════════════════════════════════════
#  App
# ═══════════════════════════════════════════════════════════════
app = FastAPI(
    title='DocScan API',
    description='Upload .docx → get PDF + per-page Markdown',
    version='1.0',
    docs_url=None,          # 关闭 Swagger UI：公网不暴露交互式文档
    # openapi.json 仍生成（status.sh 用），但见下方 _PUBLIC_API_PATHS：它也要 key
)

# Public paths that skip the API key. Only /api/health (polled by start.sh).
# /openapi.json and the (now-disabled) Swagger UI require the key — don't leak
# the endpoint schema to anonymous callers on a public deploy.
_PUBLIC_API_PATHS = ('/api/health',)


def _bearer(header: str) -> str:
    parts = header.split(None, 1)
    return parts[1].strip() if len(parts) == 2 and parts[0].lower() == 'bearer' else ''


@app.middleware('http')
async def _require_api_key(request: Request, call_next):
    """Gate every /api/* route behind DOCSCAN_API_KEY (X-API-Key header or an
    Authorization: Bearer token). Frontend pages and the ONLYOFFICE reverse
    proxy are intentionally left open. Registered before the CORS middleware so
    the latter sits outermost and answers OPTIONS preflights without a key."""
    path = request.url.path
    # 保护 /api/* 与 /openapi.json（后者不在 /api/ 前缀下，需单独覆盖）
    if (path.startswith('/api/') or path.startswith('/openapi')) and not path.startswith(_PUBLIC_API_PATHS):
        provided = request.headers.get('x-api-key') or _bearer(request.headers.get('authorization', ''))
        if not provided or provided not in _valid_keys():
            return JSONResponse({'detail': 'invalid or missing API key'}, status_code=401)
        request.state.api_key = provided   # 供下游访问控制（S6）使用
    return await call_next(request)

# ——— per-IP rate limiting (sliding window). Registered before CORS so CORS
# sits outermost (answers preflight); rate_limit wraps the key check. ———
_rate_store = {}   # {ip: {'read': [...ts], 'write': [...ts]}}

@app.middleware('http')
async def _rate_limit(request: Request, call_next):
    """Writes (POST/PUT/PATCH) are capped tighter than reads, and read/write
    quotas are tracked in SEPARATE buckets so GET traffic (health polls,
    get_md fetches) never burns the write budget. Behind a reverse proxy set
    DOCSCAN_TRUSTED_PROXY=1 so we trust X-Forwarded-For."""
    ip = request.client.host if request.client else ''
    if os.environ.get('DOCSCAN_TRUSTED_PROXY'):
        xff = request.headers.get('x-forwarded-for', '')
        if xff:
            ip = xff.split(',')[0].strip()
    if ip:
        now = time.monotonic()
        kind = 'write' if request.method in ('POST', 'PUT', 'PATCH') else 'read'
        buckets = _rate_store.setdefault(ip, {'read': [], 'write': []})
        win = [t for t in buckets[kind] if now - t < RATE_WINDOW]
        limit = RATE_WRITE if kind == 'write' else RATE_READ
        if len(win) >= limit:
            return JSONResponse({'detail': 'rate limit exceeded, retry later'}, status_code=429)
        win.append(now)
        buckets[kind] = win
        if len(_rate_store) > 50000:   # guard against unbounded IP-table growth
            _rate_store.clear()
    return await call_next(request)

# Added *after* the key + rate-limit middlewares so CORS sits outermost and
# answers OPTIONS preflights without consuming a rate-limit slot.
app.add_middleware(CORSMiddleware, allow_origins=CORS_ORIGINS, allow_methods=['*'], allow_headers=['*'])

# ——— helper ———
def _store(docx_name, pdf_path, pages, page_count):
    fid = uuid.uuid4().hex[:10]
    dest = PDFS_DIR / f'{fid}.pdf'
    shutil.move(str(pdf_path), str(dest))
    (MDS_DIR / f'{fid}.json').write_text(json.dumps(pages, ensure_ascii=False), 'utf-8')
    meta = dict(id=fid, fileName=docx_name, totalPages=page_count,
                pdfUrl=f'/api/pdf/{fid}', mdUrl=f'/api/md/{fid}', pages=pages,
                created=time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime()))
    conversions[fid] = meta
    return meta

# ═══════════════════════════════════════════════════════════════
#  API endpoints
# ═══════════════════════════════════════════════════════════════

@app.get('/api/health')
def health(): return dict(status='ok', service='DocScan API', version='1.0.0')

@app.post('/api/convert')
async def convert(request: Request, file: UploadFile = File(description='.docx file')):
    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(400, 'Only .docx accepted')
    base = uuid.uuid4().hex[:12]
    dx = DOCS_DIR / f'{base}.docx'
    pf = PDFS_DIR / f'{base}.pdf'
    dx.write_bytes(_validate_docx(await _read_capped(file)))
    loop = asyncio.get_running_loop()
    try:
        async with _convert_slot():
            n = await loop.run_in_executor(None, _convert_docx_to_pdf, dx, pf)
            pages = await loop.run_in_executor(None, _extract_pdf_pages, pf)
        meta = _store(file.filename, pf, pages, n or len(pages))
        _record_owner(meta['id'], request.state.api_key)
        dx.unlink(missing_ok=True)
        return JSONResponse(meta)
    except Exception as e:
        dx.unlink(missing_ok=True); pf.unlink(missing_ok=True)
        log.exception('docx→pdf conversion failed')
        raise HTTPException(500, 'conversion failed (see server logs)')

@app.get('/api/pdf/{fid}')
def pdf(fid: str, request: Request):
    _assert_owner(fid, request)
    p = _safe_path(PDFS_DIR, fid, 'pdf')
    if not p.exists(): raise HTTPException(404, 'not found')
    return FileResponse(str(p), media_type='application/pdf', filename=f'{fid}.pdf')

@app.get('/api/md/{fid}')
def md_all(fid: str, request: Request):
    _assert_owner(fid, request)
    j = _safe_path(MDS_DIR, fid, 'json')
    if not j.exists(): raise HTTPException(404, 'not found')
    pages = json.loads(j.read_text('utf-8'))
    return dict(id=fid, totalPages=len(pages), pages=pages,
                fileName=conversions.get(fid, {}).get('fileName',''))

@app.get('/api/md/{fid}/{page}')
def md_page(fid: str, page: int, request: Request):
    _assert_owner(fid, request)
    j = _safe_path(MDS_DIR, fid, 'json')
    if not j.exists(): raise HTTPException(404, 'not found')
    pages = json.loads(j.read_text('utf-8'))
    if page < 1 or page > len(pages): raise HTTPException(404, f'page {page} out of range')
    return dict(id=fid, page=page, totalPages=len(pages), markdown=pages[page-1])

@app.get('/api/conversions')
def list_conv(request: Request):
    """仅返回当前 key 归属的转换记录——多 key 场景下不互相可见
    （修复此前整表泄露 id/fileName 元数据的问题）。"""
    key = getattr(request.state, 'api_key', None)
    if not key:
        return list(conversions.values())   # 中间件已强制 key，此处仅兜底
    return [m for m in conversions.values() if _owner_of(m.get('id')) == key]

# ═══════════════════════════════════════════════════════════════
#  md → docx, and docx placeholder/crossref editing
# ═══════════════════════════════════════════════════════════════

def _docx_meta(fid, file_name):
    return dict(id=fid, fileName=file_name, docxUrl=f'/api/docx/{fid}',
                created=time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime()))

def _docx_path(fid):
    p = _safe_path(DOCX_DIR, fid, 'docx')
    if not p.exists():
        raise HTTPException(404, 'not found')
    return p

def _md_to_docx_sync(md_path, docx_path):
    """pandoc + docx post-processing. Synchronous and blocking, so callers must
    offload it to a worker thread (see md2docx below)."""
    r = subprocess.run(['pandoc', str(md_path), '-o', str(docx_path), '--standalone'],
                       capture_output=True, text=True, timeout=60)
    if r.returncode != 0:
        raise RuntimeError(r.stderr.strip() or 'pandoc exited non-zero')
    doc = Document(str(docx_path))
    docx_ops.convert_hr_to_page_breaks(doc)
    docx_ops.autofit_tables(doc)
    doc.save(str(docx_path))


@app.post('/api/md2docx')
async def md2docx(request: Request, file: UploadFile = File(description='.md file')):
    if not file.filename or not file.filename.lower().endswith('.md'):
        raise HTTPException(400, 'Only .md accepted')
    fid = uuid.uuid4().hex[:10]
    md_path = DOCS_DIR / f'{fid}.md'
    docx_path = DOCX_DIR / f'{fid}.docx'
    md_path.write_bytes(_validate_md(await _read_capped(file)))
    loop = asyncio.get_running_loop()
    try:
        # pandoc + python-docx are blocking — run them in the thread pool so they
        # don't freeze uvicorn's event loop (and /api/health) for up to the
        # pandoc timeout. Shares the convert slot (concurrency cap + queue cap).
        async with _convert_slot():
            await loop.run_in_executor(None, _md_to_docx_sync, md_path, docx_path)
    except Exception as e:
        md_path.unlink(missing_ok=True)
        docx_path.unlink(missing_ok=True)
        log.exception('md2docx failed')
        raise HTTPException(500, 'md2docx failed (see server logs)')
    md_path.unlink(missing_ok=True)
    _record_owner(fid, request.state.api_key)
    meta = _docx_meta(fid, file.filename)
    docx_docs[fid] = meta
    return JSONResponse(meta)

@app.post('/api/docx/upload')
async def upload_docx(request: Request, file: UploadFile = File(description='.docx file')):
    """Accept an existing .docx (e.g. from an external generator like
    generate_docx.js) and register it for editing — placeholder listing /
    replacement, cross-reference insertion, preview, and download — exactly
    as if it had been created by /api/md2docx.

    No conversion is performed; the file is stored as-is.
    """
    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(400, 'Only .docx accepted')
    fid = uuid.uuid4().hex[:10]
    docx_path = DOCX_DIR / f'{fid}.docx'
    docx_path.write_bytes(_validate_docx(await _read_capped(file)))
    _record_owner(fid, request.state.api_key)
    meta = _docx_meta(fid, file.filename)
    docx_docs[fid] = meta
    return JSONResponse(meta)

@app.get('/api/docx/{fid}')
def get_docx(fid: str, request: Request):
    _assert_owner(fid, request)
    p = _docx_path(fid)
    return FileResponse(str(p), media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         filename=f'{fid}.docx', headers={'Cache-Control': 'no-store'})

@app.get('/api/docx/{fid}/placeholders')
def get_placeholders(fid: str, request: Request):
    _assert_owner(fid, request)
    p = _docx_path(fid)
    doc = Document(str(p))
    placeholders = [ph.to_dict() for ph in docx_ops.list_placeholders(doc)]
    return dict(id=fid, count=len(placeholders), placeholders=placeholders)

class ReplaceRequest(BaseModel):
    replacements: dict[str, str]   # {placeholder_id: new_text}

@app.post('/api/docx/{fid}/replace')
def replace_placeholders(fid: str, body: ReplaceRequest, request: Request):
    _assert_owner(fid, request)
    p = _docx_path(fid)
    doc = Document(str(p))
    count = docx_ops.replace_placeholders(doc, body.replacements)
    doc.save(str(p))
    return dict(id=fid, replaced=count)

@app.get('/api/docx/{fid}/tables')
def get_tables(fid: str, request: Request):
    _assert_owner(fid, request)
    p = _docx_path(fid)
    doc = Document(str(p))
    return dict(id=fid, tables=docx_ops.list_tables(doc))

@app.get('/api/docx/{fid}/preview')
def get_preview(fid: str, request: Request):
    _assert_owner(fid, request)
    """Lightweight full-text preview of the current docx — body paragraphs
    (the only pool eligible as a crossref keyword source) plus all tables.
    Pure local read, no ONLYOFFICE round-trip, so it's fast enough to call
    after every edit to show the effect immediately.
    """
    p = _docx_path(fid)
    doc = Document(str(p))
    return dict(id=fid,
                paragraphs=docx_ops.list_body_paragraphs(doc),
                tables=docx_ops.list_tables(doc))

class CrossrefRequest(BaseModel):
    keyword: str                    # exact text to locate in the document body
    cellPath: str                   # e.g. "table[13].row[1].cell[2]" — from GET .../tables
    paragraphPath: str | None = None  # e.g. "paragraph[5]", from GET .../preview — required
                                       # when `keyword` occurs in more than one body paragraph

@app.post('/api/docx/{fid}/crossref')
async def add_crossref(fid: str, body: CrossrefRequest, request: Request):
    _assert_owner(fid, request)
    p = _docx_path(fid)
    doc = Document(str(p))
    try:
        bookmark = docx_ops.add_page_crossref(doc, body.keyword, body.cellPath, body.paragraphPath)
    except ValueError as e:
        raise HTTPException(400, str(e))
    doc.save(str(p))

    # Bake the real page number into the field via ONLYOFFICE, then swap
    # the recalculated file back in as the canonical stored docx.
    loop = asyncio.get_running_loop()
    recalced = DOCX_DIR / f'{fid}-recalc.docx'
    try:
        async with _convert_slot():
            await loop.run_in_executor(None, _recalculate_fields_docx, p, recalced)
        shutil.move(str(recalced), str(p))
    except Exception as e:
        recalced.unlink(missing_ok=True)
        log.exception('page recalculation failed')
        raise HTTPException(500, 'page recalculation failed (see server logs)')
    return dict(id=fid, bookmark=bookmark, cellPath=body.cellPath)


class CrossrefBatchItem(BaseModel):
    keyword: str
    cellPath: str
    paragraphPath: str | None = None


class CrossrefBatchRequest(BaseModel):
    items: list[CrossrefBatchItem]
    recalc: bool | None = None            # 默认 True:插入完成后做一次 ONLYOFFICE 重算
    continueOnError: bool | None = None   # 默认 True:单条失败不中断批次


def _crossref_batch_sync(p, items, continue_on_error):
    """线程池里跑:打开 doc → 批量插入 → 保存。不重算。打开/保存失败会抛
    (批次级致命,由路由转 500)。整段同步 CPU 工作必须脱离事件循环。"""
    doc = Document(str(p))
    results = docx_ops.add_page_crossref_batch(doc, items, continue_on_error)
    doc.save(str(p))
    return results


@app.post('/api/docx/{fid}/crossref_batch')
async def add_crossref_batch(fid: str, body: CrossrefBatchRequest, request: Request):
    _assert_owner(fid, request)
    p = _docx_path(fid)
    if not body.items:                    # 空批次不重算
        return dict(id=fid, total=0, succeeded=0, failed=0, recalc='skipped', items=[])

    do_recalc = True if body.recalc is None else body.recalc
    coe = True if body.continueOnError is None else body.continueOnError

    loop = asyncio.get_running_loop()
    results = await loop.run_in_executor(           # 整段 CPU 工作脱离事件循环(硬性)
        None, _crossref_batch_sync, p,
        [it.model_dump() for it in body.items], coe)

    succeeded = sum(1 for r in results if r.get('status') == 'ok')
    recalc_status = 'skipped'
    if do_recalc and succeeded > 0:                 # 有成功项才需重算
        recalced = DOCX_DIR / f'{fid}-recalc.docx'
        try:
            async with _convert_slot():
                await loop.run_in_executor(None, _recalculate_fields_docx, p, recalced)
            shutil.move(str(recalced), str(p))
            recalc_status = 'ok'
        except Exception:
            recalced.unlink(missing_ok=True)
            log.exception('batch page recalculation failed')
            recalc_status = 'failed'                # 不抛 500:保留已插入域,交调用方决策
    return dict(id=fid, total=len(results), succeeded=succeeded,
                failed=len(results) - succeeded, recalc=recalc_status, items=results)


# ═══════════════════════════════════════════════════════════════
#  Apply styles from a sample docx (heading + body) onto a stored docx
#  — reuses the same fid/owner model as replace/crossref.
# ═══════════════════════════════════════════════════════════════

@app.post('/api/docx/{fid}/apply-style')
async def apply_style(fid: str, request: Request,
                      sample: UploadFile = File(description='样式样本 .docx')):
    """Transfer heading + body paragraph styles from `sample` onto the stored
    target docx (fid), rewriting it in place so subsequent GETs reflect the
    new look. Writes a {fid}.docx.bak backup first (rollback on failure).

    Pairing is role-based: headings pair by outlineLvl (0..N) across the two
    documents regardless of styleId/name drift; body pairs by semantic name.
    Theme fonts (majorFont/minorFont) are synced so themeFont refs resolve.
    """
    _assert_owner(fid, request)
    p = _docx_path(fid)
    if not sample.filename or not sample.filename.lower().endswith('.docx'):
        raise HTTPException(400, 'Only .docx accepted as sample')
    sample_bytes = _validate_docx(await _read_capped(sample))
    sp = DOCX_DIR / f'{fid}-sample.docx'
    sp.write_bytes(sample_bytes)

    # Snapshot the *original* target before mutating, so we can roll back on
    # failure and the user can recover the pre-style file. Only the first
    # application creates the .bak — repeat applies keep the oldest original.
    bak = p.with_suffix('.docx.bak')
    created_bak = False
    if not bak.exists():
        shutil.copy2(str(p), str(bak))
        created_bak = True
    try:
        loop = asyncio.get_running_loop()
        # python-docx load/mutate/save + zip theme rewrite are blocking.
        result = await loop.run_in_executor(None, style_ops.apply_sample_styles, p, sp)
    except Exception as e:
        if created_bak:                      # restore the untouched original
            shutil.move(str(bak), str(p))
        log.exception('apply-style failed')
        raise HTTPException(500, 'style application failed (see server logs)')
    finally:
        sp.unlink(missing_ok=True)
    meta = _docx_meta(fid, p.name)
    docx_docs[fid] = meta
    return JSONResponse(dict(id=fid, fileName=p.name, docxUrl=f'/api/docx/{fid}',
                             applied=result['applied'],
                             themeFontsSynced=result['themeFontsSynced'],
                             docDefaultsSynced=result['docDefaultsSynced'],
                             numberingSynced=result['numberingSynced']))

# ═══════════════════════════════════════════════════════════════
#  Frontend demo  ( / → index.html )
# ═══════════════════════════════════════════════════════════════

@app.get('/')
@app.get('/index.html')
def frontend():
    html = (ROOT / 'index.html').read_text('utf-8')
    return HTMLResponse(html, headers={'Cache-Control':'no-cache'})

@app.get('/edit.html')
def edit_frontend():
    html = (ROOT / 'edit.html').read_text('utf-8')
    return HTMLResponse(html, headers={'Cache-Control':'no-cache'})

# ═══════════════════════════════════════════════════════════════
#  ONLYOFFICE reverse proxy   (/oo/*, /coauthoring/*, etc.)
#  Keeps the viewer at the top of index.html same-origin.
# ═══════════════════════════════════════════════════════════════

PROXY_PREFIXES = ('/oo/', '/coauthoring/', '/sdkjs/', '/web-apps/',
                  '/fonts/', '/dictionaries/', '/cache/', '/doc/')

@app.api_route('/{path:path}', methods=['GET','POST','PUT','DELETE','PATCH','OPTIONS','HEAD'])
async def proxy(request: Request, path: str = ''):
    """Catch-all that forwards ONLYOFFICE resources to the Docker container."""
    full = '/' + path
    if not any(full.startswith(p) for p in PROXY_PREFIXES):
        raise HTTPException(404, f'Not found: {full}')

    # Strip /oo prefix — ONLYOFFICE backend serves from root
    backend_path = full
    if backend_path.startswith('/oo/'):
        backend_path = backend_path[3:]  # /oo/xxx → /xxx
    backend_url = OO_BACKEND + backend_path
    headers = dict(request.headers)
    # strip hop-by-hop
    for h in ('host','connection','transfer-encoding','content-length','content-encoding','accept-encoding'):
        headers.pop(h, None)
    # preserve client Host so docservice generates same-origin URLs
    headers['host'] = request.headers.get('host', 'localhost:8800')

    body = await request.body() if request.method in ('POST','PUT','PATCH') else None
    timeout = httpx.Timeout(90, connect=10)
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=False) as client:
        r = await client.request(request.method, backend_url, headers=headers, content=body)

    # rewrite headers (Location must be /-prefixed, not absolute to 8079)
    resp_headers = {}
    for k, v in r.headers.items():
        kl = k.lower()
        if kl in ('transfer-encoding','content-encoding','connection','content-length'):
            continue
        if kl == 'location':
            v = _rewrite_location(v)
        resp_headers[k] = v

    body = r.content  # read full body (httpx streaming has edge cases)
    return Response(
        content=body,
        status_code=r.status_code,
        headers=resp_headers,
        media_type=r.headers.get('content-type'),
    )

def _rewrite_location(loc: str) -> str:
    """Redirect URLs pointing back to the backend into proxy paths."""
    if loc.startswith('/'):
        return loc  # already relative
    try:
        p = urlparse(loc)
        if p.netloc in ('localhost:8079', '127.0.0.1:8079'):
            return p.path + ('?' + p.query if p.query else '')
    except Exception:
        pass
    return loc

# ═══════════════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument('--port', type=int, default=8800, help='listen port')
    p.add_argument('--host', default='0.0.0.0')
    args = p.parse_args()
    uvicorn.run('api:app', host=args.host, port=args.port, reload=False, log_level='info')
